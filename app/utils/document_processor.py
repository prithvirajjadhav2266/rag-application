import os
import json
import docx
import docx2txt
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image
import io
import base64
from typing import Dict, List, Any, Optional, Tuple
import pandas as pd
import re
from datetime import datetime

class DocumentProcessor:
    """Process Word documents and extract text, tables, and images."""
    
    def __init__(self, upload_folder: str, metadata_dir: str):
        self.upload_folder = upload_folder
        self.metadata_dir = metadata_dir
        os.makedirs(upload_folder, exist_ok=True)
        os.makedirs(metadata_dir, exist_ok=True)
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a Word document and extract all content."""
        try:
            # Extract basic text
            text_content = docx2txt.extract(file_path)
            
            # Process with python-docx for advanced features
            doc = Document(file_path)
            
            # Extract structured content
            structured_content = self._extract_structured_content(doc)
            
            # Extract metadata
            metadata = self._extract_metadata(doc, file_path)
            
            # Extract tables
            tables = self._extract_tables(doc)
            
            # Extract images
            images = self._extract_images(doc, file_path)
            
            # Create comprehensive document data
            document_data = {
                'text_content': text_content,
                'structured_content': structured_content,
                'metadata': metadata,
                'tables': tables,
                'images': images,
                'processing_timestamp': datetime.now().isoformat(),
                'file_path': file_path
            }
            
            # Save metadata
            self._save_metadata(document_data, file_path)
            
            return document_data
            
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    def _extract_structured_content(self, doc: Document) -> Dict[str, Any]:
        """Extract structured content from document."""
        content = {
            'paragraphs': [],
            'headings': [],
            'sections': []
        }
        
        current_section = {'title': 'Introduction', 'content': []}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                # Save current section if it has content
                if current_section['content']:
                    content['sections'].append(current_section)
                
                # Start new section
                current_section = {'title': text, 'content': []}
                content['headings'].append({
                    'text': text,
                    'level': para.style.name
                })
            else:
                # Regular paragraph
                content['paragraphs'].append(text)
                current_section['content'].append(text)
        
        # Add final section
        if current_section['content']:
            content['sections'].append(current_section)
        
        return content
    
    def _extract_metadata(self, doc: Document, file_path: str) -> Dict[str, Any]:
        """Extract document metadata."""
        core_props = doc.core_properties
        
        metadata = {
            'filename': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'title': core_props.title or 'Unknown',
            'author': core_props.author or 'Unknown',
            'subject': core_props.subject or 'Unknown',
            'created': core_props.created.isoformat() if core_props.created else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'section_count': len(doc.sections)
        }
        
        # Extract company information from content
        text_content = '\n'.join([p.text for p in doc.paragraphs])
        company_info = self._extract_company_info(text_content)
        metadata.update(company_info)
        
        return metadata
    
    def _extract_company_info(self, text: str) -> Dict[str, Any]:
        """Extract company-specific information from RHP document."""
        company_info = {
            'company_name': self._extract_company_name(text),
            'industry': self._extract_industry(text),
            'ipo_size': self._extract_ipo_size(text),
            'listing_date': self._extract_listing_date(text),
            'lead_managers': self._extract_lead_managers(text)
        }
        
        return company_info
    
    def _extract_company_name(self, text: str) -> Optional[str]:
        """Extract company name from RHP text."""
        # Look for common patterns in RHP documents
        patterns = [
            r'(?:company|corporation|limited|ltd\.?|inc\.?)[\s\n]*([A-Z][A-Za-z\s&]+(?:Limited|Ltd\.?|Corporation|Corp\.?|Inc\.?))',
            r'Red Herring Prospectus.*?(?:of|for)\s+([A-Z][A-Za-z\s&]+Limited)',
            r'PUBLIC ISSUE.*?([A-Z][A-Za-z\s&]+(?:LIMITED|LTD))'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _extract_industry(self, text: str) -> Optional[str]:
        """Extract industry information."""
        industry_keywords = {
            'technology': ['software', 'technology', 'IT', 'digital'],
            'pharmaceuticals': ['pharmaceutical', 'drug', 'medicine', 'healthcare'],
            'financial': ['bank', 'financial', 'insurance', 'lending'],
            'manufacturing': ['manufacturing', 'production', 'factory'],
            'retail': ['retail', 'consumer', 'shopping'],
            'real_estate': ['real estate', 'property', 'construction']
        }
        
        text_lower = text.lower()
        for industry, keywords in industry_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                return industry.replace('_', ' ').title()
        
        return None
    
    def _extract_ipo_size(self, text: str) -> Optional[str]:
        """Extract IPO size information."""
        patterns = [
            r'issue size.*?(?:rs\.?|₹)\s*(\d+(?:,\d+)*(?:\.\d+)?)\s*crore',
            r'public issue.*?(?:rs\.?|₹)\s*(\d+(?:,\d+)*(?:\.\d+)?)\s*crore',
            r'ipo.*?(?:rs\.?|₹)\s*(\d+(?:,\d+)*(?:\.\d+)?)\s*crore'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return f"₹{match.group(1)} Crores"
        
        return None
    
    def _extract_listing_date(self, text: str) -> Optional[str]:
        """Extract expected listing date."""
        patterns = [
            r'listing.*?(\d{1,2}[/-]\d{1,2}[/-]\d{4})',
            r'expected.*?list.*?(\d{1,2}[/-]\d{1,2}[/-]\d{4})',
            r'commencement.*?trading.*?(\d{1,2}[/-]\d{1,2}[/-]\d{4})'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1)
        
        return None
    
    def _extract_lead_managers(self, text: str) -> List[str]:
        """Extract lead managers information."""
        patterns = [
            r'lead manager[s]?.*?[:]\s*([A-Za-z\s,&]+)',
            r'book running lead manager[s]?.*?[:]\s*([A-Za-z\s,&]+)',
            r'brlm.*?[:]\s*([A-Za-z\s,&]+)'
        ]
        
        managers = []
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                manager_text = match.group(1)
                # Split by common separators
                manager_list = re.split(r'[,&]', manager_text)
                managers.extend([m.strip() for m in manager_list if m.strip()])
        
        return list(set(managers))  # Remove duplicates
    
    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract and process tables from document."""
        tables_data = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                'table_id': i,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'data': []
            }
            
            # Extract table data
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data['data'].append(row_data)
            
            # Try to identify table type
            table_data['type'] = self._identify_table_type(table_data['data'])
            
            tables_data.append(table_data)
        
        return tables_data
    
    def _identify_table_type(self, table_data: List[List[str]]) -> str:
        """Identify the type of table based on content."""
        if not table_data:
            return 'unknown'
        
        # Convert to text for analysis
        table_text = ' '.join([' '.join(row) for row in table_data]).lower()
        
        # Financial statement keywords
        if any(keyword in table_text for keyword in ['revenue', 'profit', 'assets', 'liabilities', 'income statement']):
            return 'financial_statement'
        elif any(keyword in table_text for keyword in ['shareholding', 'shares', 'equity']):
            return 'shareholding_pattern'
        elif any(keyword in table_text for keyword in ['risk', 'factor']):
            return 'risk_factors'
        elif any(keyword in table_text for keyword in ['price', 'band', 'issue']):
            return 'issue_details'
        else:
            return 'general'
    
    def _extract_images(self, doc: Document, file_path: str) -> List[Dict[str, Any]]:
        """Extract images from document."""
        images_data = []
        
        try:
            # This is a simplified version - actual image extraction from docx is complex
            # For a full implementation, you'd need to parse the document's XML structure
            
            # Placeholder for image extraction logic
            # In a real implementation, you'd extract images from document relationships
            images_data.append({
                'note': 'Image extraction requires additional implementation',
                'extracted_count': 0
            })
            
        except Exception as e:
            images_data.append({
                'error': f'Image extraction failed: {str(e)}',
                'extracted_count': 0
            })
        
        return images_data
    
    def _save_metadata(self, document_data: Dict[str, Any], file_path: str) -> None:
        """Save document metadata to file."""
        filename = os.path.basename(file_path)
        metadata_filename = f"{os.path.splitext(filename)[0]}_metadata.json"
        metadata_path = os.path.join(self.metadata_dir, metadata_filename)
        
        # Create a simplified version for storage
        metadata_for_storage = {
            'metadata': document_data['metadata'],
            'tables_summary': {
                'count': len(document_data['tables']),
                'types': [table['type'] for table in document_data['tables']]
            },
            'content_summary': {
                'paragraph_count': len(document_data['structured_content']['paragraphs']),
                'section_count': len(document_data['structured_content']['sections']),
                'heading_count': len(document_data['structured_content']['headings'])
            },
            'processing_timestamp': document_data['processing_timestamp']
        }
        
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata_for_storage, f, indent=2, ensure_ascii=False)
    
    def get_document_summary(self, file_path: str) -> Dict[str, Any]:
        """Get a quick summary of document content."""
        try:
            # Quick extraction
            text = docx2txt.extract(file_path)
            doc = Document(file_path)
            
            # Basic statistics
            summary = {
                'word_count': len(text.split()),
                'character_count': len(text),
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'section_count': len(doc.sections),
                'has_financial_data': self._has_financial_data(text),
                'document_type': self._identify_document_type(text)
            }
            
            return summary
            
        except Exception as e:
            return {'error': f'Failed to generate summary: {str(e)}'}
    
    def _has_financial_data(self, text: str) -> bool:
        """Check if document contains financial data."""
        financial_keywords = [
            'income statement', 'balance sheet', 'cash flow', 'revenue',
            'profit', 'assets', 'liabilities', 'equity', 'financial statement'
        ]
        
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in financial_keywords)
    
    def _identify_document_type(self, text: str) -> str:
        """Identify the type of document."""
        text_lower = text.lower()
        
        if 'red herring prospectus' in text_lower:
            return 'Red Herring Prospectus'
        elif 'draft red herring prospectus' in text_lower:
            return 'Draft Red Herring Prospectus'
        elif 'prospectus' in text_lower:
            return 'Prospectus'
        elif 'annual report' in text_lower:
            return 'Annual Report'
        else:
            return 'Unknown Document Type'